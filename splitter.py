import tkinter as tk
from tkinter import filedialog, ttk, messagebox
import os
import datetime
from docx import Document
import win32com.client
import pythoncom

class DocumentSplitter:
    def __init__(self, master, callback):
        self.master = master
        self.callback = callback
        self.filename = None
        self.title_structure = []
        self.original_title_structure = []
        self.deleted_items = set()
        
        # Configure styles
        style = ttk.Style()
        
        # Configure Treeview
        style.configure(
            "Custom.Treeview",
            rowheight=25
        )
        
        # Configure Treeview selection
        style.map(
            "Custom.Treeview",
            background=[("selected", "#007acc")],
            foreground=[("selected", "white")]
        )
        
        self.setup_ui()

    def setup_ui(self):
        # Main container frame with padding
        main_frame = ttk.Frame(self.master, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)

        # Upload section
        upload_frame = ttk.LabelFrame(main_frame, text="Document Upload", padding="10")
        upload_frame.pack(fill=tk.X, padx=5, pady=5)
        
        self.upload_button = ttk.Button(
            upload_frame, 
            text="Upload Word or PDF File",
            command=self.upload_file,
            style="Accent.TButton"
        )
        self.upload_button.pack(pady=(5, 10))

        # Depth control section
        depth_frame = ttk.LabelFrame(main_frame, text="Chapter Depth", padding="10")
        depth_frame.pack(fill=tk.X, padx=5, pady=5)

        depth_control_frame = ttk.Frame(depth_frame)
        depth_control_frame.pack(fill=tk.X, pady=5)
        
        self.depth_label = ttk.Label(depth_control_frame, text="Depth Level: 1")
        self.depth_label.pack(side=tk.LEFT, padx=(0, 10))

        self.depth_level = tk.IntVar(value=1)
        self.depth_level.trace_add('write', self.on_depth_change)
        
        self.depth_slider = ttk.Scale(
            depth_control_frame,
            from_=1,
            to=9,
            orient=tk.HORIZONTAL,
            variable=self.depth_level,
            command=self.on_scale_change
        )
        self.depth_slider.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        # Tree section
        tree_frame = ttk.LabelFrame(main_frame, text="Chapter Structure", padding="10")
        tree_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # Create tree with scrollbar
        tree_container = ttk.Frame(tree_frame)
        tree_container.pack(fill=tk.BOTH, expand=True)

        self.tree = ttk.Treeview(
            tree_container,
            selectmode="extended",
            style="Custom.Treeview"
        )
        self.tree.heading("#0", text="Title Structure")
        self.tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        scrollbar = ttk.Scrollbar(tree_container, orient=tk.VERTICAL, command=self.tree.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.tree.configure(yscrollcommand=scrollbar.set)

        # Button section
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill=tk.X, padx=5, pady=10)

        self.delete_button = ttk.Button(
            button_frame,
            text="Delete Selected Chapters",
            command=self.delete_selected
        )
        self.delete_button.pack(side=tk.LEFT, padx=5)

        self.reset_button = ttk.Button(
            button_frame,
            text="Reset All Chapters",
            command=self.reset_chapters
        )
        self.reset_button.pack(side=tk.LEFT, padx=5)

        self.split_button = ttk.Button(
            button_frame,
            text="Split Document and Save",
            command=self.split_document,
            style="Accent.TButton"
        )
        self.split_button.pack(side=tk.RIGHT, padx=5)

    def on_scale_change(self, value):
        # Convert the float value to the nearest integer
        new_value = round(float(value))
        # Set the slider to the integer value
        self.depth_level.set(new_value)

    def on_depth_change(self, *args):
        depth = self.depth_level.get()
        self.depth_label.config(text=f"Depth Level: {depth}")
        self.update_tree()

    def upload_file(self):
        filetypes = [('Word Documents', '*.docx'), ('PDF Files', '*.pdf'), ('All files', '*.*')]
        filename = filedialog.askopenfilename(title='Open a file', filetypes=filetypes)
        if filename:
            self.filename = filename
            if filename.lower().endswith('.pdf'):
                word_filename = self.convert_pdf_to_word(filename)
                if word_filename:
                    self.filename = word_filename
                else:
                    messagebox.showerror("Error", "Failed to convert PDF to Word.")
                    return
            self.analyze_document()
        else:
            messagebox.showinfo("No file selected", "Please select a file to upload.")

    def convert_pdf_to_word(self, pdf_filename):
        try:
            pythoncom.CoInitialize()
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False

            pdf_path = os.path.abspath(pdf_filename)
            doc = word.Documents.Open(pdf_path)

            word_filename = pdf_filename[:-4] + '.docx'
            doc.SaveAs2(os.path.abspath(word_filename), FileFormat=16)
            doc.Close()
            word.Quit()
            return word_filename
        except Exception as e:
            print(f"Error converting PDF to Word: {e}")
            return None
        finally:
            pythoncom.CoUninitialize()

    def analyze_document(self):
        try:
            doc = Document(self.filename)
            self.title_structure = []
            self.original_title_structure = []

            for para in doc.paragraphs:
                style = para.style.name
                text = para.text.strip()
                if text and style.startswith('Heading'):
                    try:
                        level = int(style.replace('Heading ', ''))
                    except:
                        level = 1
                    item = {'level': level, 'text': text}
                    self.title_structure.append(item)
                    self.original_title_structure.append(item.copy())

            self.deleted_items = set()
            self.update_tree()
        except Exception as e:
            print(f"Error analyzing document: {e}")
            messagebox.showerror("Error", "Failed to analyze the document.")

    def update_tree(self, event=None):
        self.tree.delete(*self.tree.get_children())
        depth = self.depth_level.get()
        parent_stack = [""]
        last_level = 0

        for item in self.title_structure:
            level = item['level']
            text = item['text']
            if level <= depth:
                while len(parent_stack) > level:
                    parent_stack.pop()
                    last_level -= 1
                parent = parent_stack[-1]
                current_node = self.tree.insert(parent, 'end', text=text)
                if len(parent_stack) <= level:
                    parent_stack.append(current_node)
                else:
                    parent_stack[level] = current_node
                last_level = level

    def get_clean_title(self, title):
        # Remove dots from any numbering at the start
        parts = title.split()
        if parts and any(c.isdigit() for c in parts[0]):
            # Remove the numbering part and keep the rest
            clean_title = ' '.join(parts[1:])
        else:
            clean_title = title
            
        # Remove any remaining dots and special characters
        clean_title = ''.join(c for c in clean_title if c.isalnum() or c.isspace())
        return clean_title

    def delete_selected(self):
        selected_items = self.tree.selection()
        if not selected_items:
            messagebox.showinfo("Info", "Please select chapters to delete")
            return

        items_to_delete = set()
        for item in selected_items:
            items_to_delete.add(self.tree.item(item)['text'])
            children = self.get_all_children(item)
            for child in children:
                items_to_delete.add(self.tree.item(child)['text'])

        self.deleted_items.update(items_to_delete)
        self.title_structure = [
            item for item in self.title_structure 
            if item['text'] not in self.deleted_items
        ]
        self.update_tree()

    def get_all_children(self, item):
        children = self.tree.get_children(item)
        result = list(children)
        for child in children:
            result.extend(self.get_all_children(child))
        return result

    def reset_chapters(self):
        self.title_structure = [item.copy() for item in self.original_title_structure]
        self.deleted_items = set()
        self.update_tree()
        messagebox.showinfo("Success", "All chapters have been restored")

    def split_document(self):
        if not self.filename:
            messagebox.showerror("Error", "Please upload a document first.")
            return

        if not self.title_structure:
            messagebox.showerror("Error", "No chapters selected for processing.")
            return

        depth = self.depth_level.get()
        try:
            doc = Document(self.filename)
            
            # Create projects directory in the app folder
            app_dir = os.path.dirname(os.path.abspath(__file__))
            projects_dir = os.path.join(app_dir, "projects")
            
            if not os.path.exists(projects_dir):
                os.makedirs(projects_dir)
            
            # Create shorter project name
            base_name = os.path.splitext(os.path.basename(self.filename))[0]
            base_name = base_name[:50]  # Increased length limit
            timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M')
            project_name = f"{base_name}_{timestamp}"
            output_folder = os.path.join(projects_dir, project_name)
            
            try:
                os.makedirs(output_folder, exist_ok=True)
            except Exception as e:
                messagebox.showerror("Error", f"Could not create project folder: {str(e)}")
                return

            sections = []
            current_doc = None
            current_level = 0
            current_title = ""

            for para in doc.paragraphs:
                style = para.style.name
                text = para.text.strip()
                if text:
                    if style.startswith('Heading'):
                        try:
                            level = int(style.replace('Heading ', ''))
                        except:
                            level = 1
                        
                        if text in {item['text'] for item in self.title_structure}:
                            if level <= depth:
                                if current_doc:
                                    sections.append({'level': current_level, 'title': current_title, 'document': current_doc})
                                current_doc = Document()
                                current_doc.add_paragraph(text, style=style)
                                current_level = level
                                current_title = text
                            else:
                                if current_doc:
                                    current_doc.add_paragraph(text, style=style)
                    else:
                        if current_doc:
                            current_doc.add_paragraph(text, style=style)

            if current_doc:
                sections.append({'level': current_level, 'title': current_title, 'document': current_doc})

            # Save the sections
            for idx, section in enumerate(sections, 1):
                try:
                    title = section['title']
                    clean_title = self.get_clean_title(title)
                    
                    # Take first few words for the filename
                    words = clean_title.split()
                    short_title = '_'.join(words[:3])  # Take only first 3 words
                    short_title = short_title[:100]  # Increased length limit
                    
                    output_filename = os.path.join(output_folder, f"{idx:02d}_{short_title}.docx")
                    
                    # Fallback to simple numbered filename if path is too long
                    if len(output_filename) > 240:
                        output_filename = os.path.join(output_folder, f"{idx:02d}.docx")
                    
                    section['document'].save(output_filename)
                    
                except Exception as e:
                    print(f"Error saving section {idx}: {str(e)}")
                    # Try fallback filename
                    try:
                        fallback_filename = os.path.join(output_folder, f"{idx:02d}.docx")
                        section['document'].save(fallback_filename)
                    except Exception as fallback_error:
                        print(f"Fallback save failed for section {idx}: {str(fallback_error)}")
                        messagebox.showerror("Error", f"Failed to save section {idx}")
                        continue

            if sections:
                messagebox.showinfo("Success", f"Document split into {len(sections)} sections and saved in:\n{output_folder}")
                self.callback(output_folder)
            else:
                messagebox.showwarning("Warning", "No sections were created during splitting.")
                
        except Exception as e:
            error_msg = f"An error occurred while splitting the document:\n{str(e)}"
            print(error_msg)
            messagebox.showerror("Error", error_msg)